import io
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.services.ingestion_service import detect_file_type
from app.services.rag_service import extract_text_from_document


def test_detecte_les_nouveaux_formats_documentaires():
    for filename in ("rapport.docx", "notes.md", "article.tex"):
        assert detect_file_type(filename) == "document"


def test_extrait_markdown_et_latex_en_utf8():
    markdown = "# Résultats\n\nLa croissance atteint 12 %."
    latex = r"\section{Résultats} La croissance atteint 12 \%."

    assert extract_text_from_document(markdown.encode("utf-8"), "notes.md") == markdown
    assert extract_text_from_document(latex.encode("utf-8"), "article.tex") == latex


def test_extrait_paragraphes_et_tableaux_docx():
    document = Document()
    document.add_heading("Rapport annuel", level=1)
    document.add_paragraph("Une synthèse du document.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Année"
    table.cell(0, 1).text = "Valeur"
    table.cell(1, 0).text = "2026"
    table.cell(1, 1).text = "42"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text_from_document(buffer.getvalue(), "rapport.docx")

    assert "Rapport annuel" in text
    assert "Une synthèse du document." in text
    assert "Année | Valeur" in text
    assert "2026 | 42" in text
